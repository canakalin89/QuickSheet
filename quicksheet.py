import streamlit as st
import os
import requests
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, Cm
import time
import google.generativeai as genai
import io
import re
import random
from datetime import date

# -----------------------------
# SAYFA YAPISI VE BAŞLANGIÇ AYARLARI
# -----------------------------
st.set_page_config(page_title="QuickSheet v1.7: 9. Sınıf Asistanı", page_icon="⚡", layout="wide")
st.title("⚡ Quicksheet v1.7: 9. Sınıf İngilizce Öğretmeni Asistanı")
st.markdown("Türkiye Yüzyılı Maarif Modeli'ne uygun, **bütünleşik ve pedagojik** materyaller üretin.")

# Session State Başlatma
if 'ai_content' not in st.session_state:
    st.session_state.ai_content = ""
if 'last_tool' not in st.session_state:
    st.session_state.last_tool = ""
if 'final_prompt' not in st.session_state:
    st.session_state.final_prompt = ""

# Gemini API anahtarı
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=API_KEY)
    model = genai.GenerativeModel("gemini-1.5-flash")
except (KeyError, AttributeError):
    st.error("Gemini API anahtarı bulunamadı veya geçersiz. Lütfen Streamlit Cloud secrets'a 'GEMINI_API_KEY' ekleyin.")
    st.stop()

# -----------------------------
# FONT YÜKLEME
# -----------------------------
@st.cache_resource
def load_font():
    font_dir = "fonts"
    if not os.path.exists(font_dir):
        os.makedirs(font_dir)
    font_files = {
        "DejaVuSans": "https://github.com/googlefonts/dejavu-fonts/raw/main/ttf/DejaVuSans.ttf",
        "DejaVuSans-Bold": "https://github.com/googlefonts/dejavu-fonts/raw/main/ttf/DejaVuSans-Bold.ttf"
    }
    try:
        for name, url in font_files.items():
            font_path = os.path.join(font_dir, f"{name}.ttf")
            if not os.path.exists(font_path):
                response = requests.get(url, timeout=20)
                response.raise_for_status()
                with open(font_path, "wb") as f:
                    f.write(response.content)
            pdfmetrics.registerFont(TTFont(name, font_path))
        return True
    except Exception as e:
        st.error(f"Font yüklenirken hata oluştu: {e}")
        return False
font_loaded = load_font()

# -----------------------------
# MÜFREDAT BİLGİSİ
# -----------------------------
meb_curriculum = {
    "9. Sınıf": {
        "Theme 1: School Life": {
            "Grammar": ["Simple Present (to be)", "Modal 'can' (possibility/ability)", "Simple Past (was/were)"],
            "Vocabulary": ["Countries and Nationalities", "Languages and Capitals", "Tourist Attractions", "School Rules"],
            "Reading": ["A Student's First Day at School (Blog Post)", "School Club Descriptions", "An E-mail from an Exchange Student", "A Dialogue about School Rules", "A Short Text about a Famous Landmark", "A Country Profile (Fact File)", "Travel Forum Posts", "A Quiz about World Capitals", "An Interview with a Foreign Student", "A School Trip Itinerary"],
            "Speaking": ["Introducing Yourself (Role-Play)", "Describing Your Country to a Tourist", "Asking for and Giving Directions in School", "Planning a School Trip with a Partner", "Discussing School Rules", "A Short Presentation about a Famous Landmark", "Comparing Two Schools", "Role-Play: Meeting a New Exchange Student", "A Mini-Debate: School Uniforms Yes or No?", "Giving a Short 'Welcome to Our School' Speech"],
            "Writing": ["Writing a Welcome E-mail to a New Student", "Paragraph Describing a Dream Holiday Destination", "Writing a Postcard from Abroad", "Creating a School Club Poster", "A Short Blog Post about Your School", "A 'Top 5' List of Places to Visit in Your City", "A Short Dialogue between Two New Students", "Answering Questions about Your Country", "Writing a Short Quiz about Nationalities", "Planning a Weekend Trip for a Visitor"],
            "Pronunciation": ["Long and short vowels (/æ/, /ɑː/, /eɪ/)", "Consonants (/b/, /k/, /d/)"]
        },
        "Theme 2: Classroom Life": {
            "Grammar": ["Simple Present (routines)", "Adverbs of Frequency", "Imperatives"],
            "Vocabulary": ["Daily Routines", "Study Habits", "Classroom Objects", "Instructions"],
            "Reading": ["An Interview with a Successful Student", "Effective Study Tips (Listicle)", "Comparing Two Students' Daily Schedules", "A Diary Entry about a Busy Day", "A Survey about Classroom Habits", "A Set of Rules for the School Library", "A 'How-To' Guide for Taking Good Notes", "An Article about Different Learning Styles", "A Text Message Conversation about Homework", "A Profile of a Favorite Teacher"],
            "Speaking": ["Describing Your Typical School Day", "Giving Instructions to a Classmate", "Role-Play: Asking a Teacher for Help", "Discussing Good and Bad Study Habits", "Comparing Your Weekend Routine with a Partner's", "A Short Talk about Your Morning Routine", "Explaining How to Do Something", "Group Discussion: How to Make Our Classroom Better", "Role-Play: Forgetting Your Homework", "A Mini-Presentation on an Effective Study Tip"],
            "Writing": ["Creating a Weekly Study Plan", "Paragraph Describing a Perfect School Day", "An E-mail Giving Study Advice to a Friend", "A Diary Entry about Today's Lessons", "A Set of Instructions for a Classroom Game", "A 'Top 5' List of Good Study Habits", "A Short Paragraph about Your Morning Routine", "A Note to a Teacher Explaining an Absence", "Creating a 'Classroom Rules' Poster", "A Short Story about a Funny Classroom Moment"],
            "Pronunciation": ["Vowels (/e/, /æ/)", "Consonants (/f/, /g/, /dʒ/, /h/)"]
        },
        "Theme 3: Personal Life": {
            "Grammar": ["Adjectives (personality/appearance)", "'too' / 'enough'"],
            "Vocabulary": ["Physical Appearance", "Personality Traits", "Hobbies and Interests"],
            "Reading": ["A Magazine Article about a Celebrity", "A Fictional Character Profile", "An Online Forum Discussion about Hobbies", "A 'Guess Who?' Description Game", "A Blog Post about a New Hobby", "A Comparison of Two Friends", "A Short Story about a Shy Character", "A Survey about Personality Types", "An Advice Column about Making Friends", "A Review of a Mobile App for Hobbies"],
            "Speaking": ["Describing a Friend's Personality and Appearance", "Talking about Your Hobbies and Interests", "Role-Play: Complimenting a Friend", "Guessing Game: Describing a Mystery Person", "Comparing Your Hobbies with a Partner's", "A Short Talk about Your Favorite Fictional Character", "Discussing Personality Traits for a Job", "Role-Play: Inviting a Friend to Join a Hobby", "Explaining Why You Like a Certain Hobby", "A Mini-Debate: Indoor vs. Outdoor Hobbies"],
            "Writing": ["Paragraph Describing a Best Friend", "Creating a Fictional Character Profile", "Writing a Blog Post about a Favorite Hobby", "A 'Wanted' Poster for a Lost Pet (with description)", "A Short Paragraph about Your Personality", "Describing a Family Member's Appearance and Personality", "A 'Top 3' List of Your Favorite Hobbies", "An Online Profile for a Social Media Site", "A Short Story Featuring a Brave Character", "An Email to a Friend Suggesting a New Hobby"],
            "Pronunciation": ["Vowels (/iː/, /ɪ/, /aɪ/)", "Consonants (/ʒ/, /k/, /l/)"]
        },
        "Theme 4: Family Life": {
            "Grammar": ["Simple Present (jobs and work routines)", "Prepositions of place (in/on/at for workplaces)"],
            "Vocabulary": ["Family members", "Jobs and Occupations", "Workplaces", "Work Activities"],
            "Reading": ["A 'Day in the Life' Article about a Doctor", "Informational Texts about Different Jobs", "An Interview with a Parent about Their Job", "A Family Tree with Job Descriptions", "A Short Story about a Family Business", "Comparing Office Work vs. Outdoor Work", "A Job Advertisement", "A Career Advice Blog Post", "A Text about Unusual Jobs", "A Biography of a Successful Entrepreneur"],
            "Speaking": ["Talking about your Family Members' Occupations", "Asking a Friend about Their Parents' Jobs", "Role-Play: A Job Interview (Simple)", "A Short Presentation about a Dream Job", "Discussing the Pros and Cons of a Job", "Comparing Your Parents' Jobs", "Describing a Workplace", "Explaining What a Person Does in a Specific Job", "A Mini-Debate: Is It Better to Work Indoors or Outdoors?", "Role-Play: Career Day at School"],
            "Writing": ["Paragraph Describing a Family Member's Job", "Describing a Dream Job", "A Simple Job Application E-mail", "A 'Bring Your Parent to School Day' Invitation", "A Short Paragraph about 'My Future Career'", "Creating a 'Job Profile' for a Family Member", "A List of Questions to Interview a Professional", "A Short Story about a Child Helping a Parent at Work", "A Thank-You Note to a Working Parent", "Describing a Workplace"],
            "Pronunciation": ["Vowels (/ɒ/, /ɔː/)", "Consonants (/m/, /n/, /p/)"]
        },
        "Theme 5: House & Neighbourhood": {
            "Grammar": ["Present Continuous", "There is/are", "Possessive adjectives"],
            "Vocabulary": ["Types of houses", "Rooms", "Furniture", "Household chores"],
            "Reading": ["Real Estate Ads (House Descriptions)", "Informational Text on Unique House Types", "A Blog Post about Life in a Neighbourhood", "A Dialogue between New Neighbours", "A 'House Tour' Article from a Magazine", "A Short Story Set in a Mysterious House", "A Survey about Household Chores", "An Article Comparing City and Suburb Life", "A Set of Rules for an Apartment Building", "A Review of a Local Park or Cafe"],
            "Speaking": ["Describing Your Own House in Detail", "Describing a Picture of a Room (What People Are Doing)", "Role-Play: Calling a Landlord to Report a Problem", "Comparing Your House with a Friend's", "A Short Presentation about Your Dream House", "Giving a Tour of Your Imaginary House", "Discussing Pros and Cons of a Flat vs. a House", "Describing Your Neighbourhood to a Visitor", "Role-Play: Asking a Neighbour for a Favor", "Explaining How to Do a Household Chore"],
            "Writing": ["A Paragraph Describing a Dream House", "Writing a 'For Rent' Advertisement", "A Note to a Neighbour about a Party", "A 'Welcome to the Neighbourhood' Card for a New Family", "A List of Pros and Cons of Your Neighbourhood", "A Short Paragraph about Your Favorite Room", "Describing What is Happening in a Picture of a House", "An E-mail to a Friend Describing Your New Home", "Creating a List of House Rules for a Guest", "A Short Story about a Lost Item in the House"],
            "Pronunciation": ["Consonants (/q/ as in quick, /r/, /s/, /ʃ/ as in shower)"]
        },
        "Theme 6: City & Country": {
            "Grammar": ["Present Simple vs. Present Continuous", "Wh- questions", "'or' for options"],
            "Vocabulary": ["Food culture", "Food festivals", "Ingredients", "Local and international dishes"],
            "Reading": ["A Newspaper Article about a Food Festival", "The History of a Traditional Turkish Dish", "A Chef's Blog Post", "A Restaurant Review", "A Travel Guide Excerpt about Local Cuisine", "An Interview with a Street Food Vendor", "A Dialogue about Ordering Food", "A Recipe for a Simple Dish", "A Comparison of Two Different Cuisines", "A 'Top 5' List of Must-Try Foods in a City"],
            "Speaking": ["Role-Playing a Scene at a Food Festival", "Asking for Options in a Restaurant", "Describing Your Favorite Food", "Comparing Street Food and Restaurant Food", "A Short Presentation on a Traditional Dish", "Discussing the Importance of Food Culture", "Role-Play: A Phone Call to Order a Pizza", "Explaining a Simple Recipe Step-by-Step", "A Mini-Debate: Is Fast Food Always Unhealthy?", "Interviewing a Partner about Their Eating Habits"],
            "Writing": ["Writing a Blog Post about a Food Festival", "Writing a Recipe for a Favorite Dish", "A Simple Restaurant Review", "An E-mail to a Friend Recommending a Restaurant", "A Paragraph Comparing Two Types of Food", "Creating a Menu for a Small Cafe", "A Short Story about a Memorable Meal", "A 'Food Diary' for a Day", "Describing Your Country's National Dish", "A Thank-You Note to a Host for a Delicious Meal"],
            "Pronunciation": ["Vowels (/uː/, /ʊ/)", "Consonants (/t/, /ð/, /θ/, /v/)"]
        },
        "Theme 7: World & Nature": {
            "Grammar": ["Simple Past (was/were, there was/were)", "Modal 'should' (advice)"],
            "Vocabulary": ["Endangered animals", "Habitats", "Environmental problems (habitat loss, pollution)"],
            "Reading": ["A Fact File about an Endangered Animal", "A News Report on a Conservation Project", "A Volunteer's Diary from a Wildlife Shelter", "An Article about the Dangers of Plastic Pollution", "A Short Story about Saving an Animal", "An Interview with a Marine Biologist", "A 'Then and Now' Text about a Habitat", "A List of 'Do's and Don'ts' for Protecting Nature", "A Poem about the Beauty of Nature", "A Debate Text about Zoos"],
            "Speaking": ["Giving Advice on How to Protect Nature", "Discussing Local Environmental Problems", "A Short Presentation on an Endangered Animal", "Role-Play: Convincing Someone to Recycle", "Comparing Two Environmental Problems", "Describing a Beautiful Natural Place", "A Mini-Debate: Should Zoos Exist?", "Explaining the Importance of a Habitat", "Role-Play: A News Report about a Successful Conservation Project", "Brainstorming Solutions for Plastic Pollution in a Group"],
            "Writing": ["Paragraph on How to Protect Nature", "Writing an E-mail to an Animal Protection Organization", "Text for a 'Save the Planet' Poster", "A Short Paragraph Giving Advice to a Friend on Being Eco-Friendly", "A Fact File about an Endangered Animal from Your Country", "A Short Story from an Animal's Perspective", "A 'Promise to the Planet' Pledge", "A Letter to the Editor about a Local Environmental Issue", "Creating a Slogan for a Conservation Campaign", "Describing a Beautiful Natural Place You Have Visited"],
            "Pronunciation": ["Diphthongs (/eə/ as in bear, /ɪə/ as in deer)", "Consonants (/w/, /ks/)"]
        },
        "Theme 8: Universe & Future": {
            "Grammar": ["Future Simple (will) for predictions and beliefs", "Simple Present for describing films"],
            "Vocabulary": ["Films and film genres", "Futuristic ideas (robots, space exploration)", "Technology"],
            "Reading": ["A Sci-Fi Movie Review", "An Article about Future Technologies", "An Interview with a Famous Director", "A Short Sci-Fi Story", "A 'Top 5' List of Predictions for the Future", "A Blog Post about the Impact of Technology on Daily Life", "A Comparison of Two Sci-Fi Films", "A Profile of a Tech Inventor", "A Text about the Possibility of Life on other Planets", "A Humorous Look at Old Predictions about the Future"],
            "Speaking": ["Making Predictions about the Future (e.g., schools, transport)", "Talking about Your Favorite Films and Genres", "A Short Presentation on a Piece of Future Technology", "Role-Play: Explaining Today's Technology to Someone from the Past", "A Mini-Debate: Will Robots Be Good for Humanity?", "Comparing Two Movie Characters", "Describing a Sci-Fi Movie Plot to a Friend", "Discussing the Pros and Cons of Social Media", "Explaining How a Piece of Technology Works", "Role-Play: Pitching a New Sci-Fi Movie Idea"],
            "Writing": ["A Day in 2050 (Short Story)", "Paragraph Describing a Favorite Sci-Fi Movie", "Describing a Dream Robot", "An E-mail to a Friend about a Movie You Just Watched", "A 'For and Against' Paragraph about Artificial Intelligence", "My Predictions for the Next 10 Years (List)", "Creating a Character for a Sci-Fi Story", "A Short Review of a Tech Gadget", "A Story about Time Travel", "Writing the Plot Summary for a New Movie Idea"],
            "Pronunciation": ["Diphthong (/əʊ/ as in show)", "Consonants (/j/ as in year, /z/)"]
        }
    }
}


# -----------------------------
# SIDEBAR - KULLANICI ARAYÜZÜ
# -----------------------------
with st.sidebar:
    st.header("1. Adım: Seçimlerinizi Yapın")
    selected_grade = "9. Sınıf"
    units = list(meb_curriculum[selected_grade].keys())
    selected_unit = st.selectbox("Ünite Seçin", units)
    
    selected_tool = st.radio(
        "Materyal Türü Seçin",
        ["Günlük Ders Planı", "Çalışma Sayfası", "Ders Planı (Genel)", "Dinleme Aktivitesi Senaryosu", "Ünite Tekrar Testi", "Değerlendirme Rubriği", "Ek Çalışma (Farklılaştırılmış)"],
        captions=["Belirli bir konu için 40dk'lık plan", "Alıştırma kağıdı üretir", "Ünite geneli için ders planı", "Dinleme metni ve soruları yazar", "Üniteyi özetleyen test hazırlar", "Puanlama anahtarı oluşturur", "Destekleyici veya ileri düzey aktivite"]
    )
    
    plan_language = "İngilizce"
    if selected_tool in ["Günlük Ders Planı", "Ders Planı (Genel)"]:
        plan_language = st.selectbox("Plan Dili", ["Türkçe", "İngilizce"])

    if selected_tool == "Günlük Ders Planı":
        st.date_input("Ders Tarihi", value=date.today(), key="lesson_date")
        st.text_input("Bugünkü Dersin Konusu / Odak Noktası", placeholder="Örn: Kelime tekrarı, Sayfa 28, Alıştırma 3", key="lesson_context")

    skill_needed = selected_tool in ["Çalışma Sayfası", "Değerlendirme Rubriği", "Ek Çalışma (Farklılaştırılmış)", "Ünite Tekrar Testi"]
    if skill_needed:
        skill_options = list(meb_curriculum[selected_grade][selected_unit].keys())
        selected_skill = st.selectbox("Odaklanılacak Beceriyi Seçin", skill_options)
        
        topic_options = meb_curriculum[selected_grade][selected_unit].get(selected_skill, [])
        if topic_options:
            selected_topics_from_list = st.multiselect(
                "Hangi Konulara Odaklanılsın?",
                topic_options,
                default=topic_options[0] if topic_options else []
            )
        else:
            selected_topics_from_list = []
        
        custom_topics = st.text_input("Veya Kendi Konularınızı Ekleyin (virgülle ayırarak)", key="custom_topics_input")
            
    else:
        selected_skill = "Genel"
        selected_topics_from_list = []
        custom_topics = ""


    with st.expander("Gelişmiş Ayarlar"):
        include_clil = st.checkbox("CLIL İçeriği Ekle (Teknoloji, Bilim vb.)")
        include_reflection = st.checkbox("Yansıtma Aktivitesi Ekle (Öz değerlendirme)")
        
        if selected_tool in ["Çalışma Sayfası", "Ünite Tekrar Testi"]:
            num_questions = st.slider("Soru Sayısı", 3, 10, 6)
        
        if selected_tool == "Ek Çalışma (Farklılaştırılmış)":
            differentiation_type = st.radio("Aktivite Düzeyi", ["Destekleyici (Supporting)", "İleri Düzey (Expansion)"])

# -----------------------------
# PROMPT OLUŞTURMA FONKSİYONLARI (GÜNCELLENDİ)
# -----------------------------
def get_activity_suggestions(skill):
    suggestions = {
        "Grammar": ["a sentence transformation task", "a correct-the-mistake exercise", "a sentence building task from jumbled words"],
        "Vocabulary": ["a matching exercise (word to definition)", "a fill-in-the-blanks story using a word bank", "an odd-one-out task"],
        "Reading": ["a short text with True/False questions", "a text where students match paragraphs to headings"],
        "Speaking": ["a role-play scenario", "a set of discussion questions", "a picture description task"],
        "Writing": ["a guided paragraph writing task with prompts", "an email writing task based on a scenario"],
        "Pronunciation": ["a minimal pairs discrimination exercise", "a tongue twister", "a 'find the sound' activity"]
    }
    return random.choice(suggestions.get(skill, ["a multiple-choice exercise"]))

def create_prompt(tool, **kwargs):
    language_instruction = f"The entire output for this plan MUST be in {kwargs.get('language')}."
    
    base_prompt = f"""
Act as an expert English teacher for the Turkish Ministry of Education's new 'Türkiye Yüzyılı Maarif Modeli'.
Create a high-quality, engaging material for a {kwargs.get('grade')} class (CEFR B1.1), aligned with the WAYMARK series.
The content must be entirely in English, unless a different language is specified for the output.
Unit: "{kwargs.get('unit')}"
"""
    
    topic_text = ""
    if kwargs.get('topics'):
        topic_text = f"**Strict Focus:** The material MUST ONLY focus on the following topic(s): {', '.join(kwargs.get('topics'))}."

    # YENİ: Becerileri entegre etme talimatı
    integration_text = ""
    skill = kwargs.get('skill')
    if skill in ["Reading", "Writing", "Speaking"]:
        unit_data = meb_curriculum[kwargs.get('grade')][kwargs.get('unit')]
        grammar_focus = ", ".join(unit_data.get("Grammar", []))
        vocab_focus = ", ".join(unit_data.get("Vocabulary", []))
        integration_text = f"""
**Pedagogical Integration Requirement:** The '{skill}' activity you create MUST naturally include and provide context for the following language points from this unit:
- **Grammar to Integrate:** {grammar_focus}
- **Vocabulary to Integrate:** {vocab_focus}
"""
    
    prompts = {
        "Günlük Ders Planı": f"""
        **Task:** Create a detailed 40-minute daily lesson plan for {kwargs.get('date')}.
        **Today's Specific Focus:** "{kwargs.get('context')}"
        {language_instruction}
        The plan MUST be well-structured, using clear headings (like **Objective:**, **Materials:**, **Warm-Up (5 min):** etc.) and bullet points.
        It must be practical and student-centered.
        """,
        "Çalışma Sayfası": f"""
        Focus Skill: "{skill}"
        {topic_text}
        {integration_text}
        Create a worksheet with exactly {kwargs.get('num_questions')} questions.
        - **Instruction for AI:** Ensure activities are diverse. Use a variety of task types like {get_activity_suggestions(skill)}.
        - **Strict Rule:** Do NOT include questions for any other skill.
        - Start with a title and a one-sentence instruction.
        - End with a separate 'Answer Key' section.
        """,
        # ... Diğer prompt'lar da integration_text'i kullanabilir ...
        "Ders Planı (Genel)": f"""
        Create a 40-minute lesson plan for the selected unit.
        {language_instruction}
        - The plan must be well-structured with clear headings and bullet points.
        - Include: Objective, Key Language, Materials.
        - Structure in three stages: Warm-Up/Lead-In (5-10 min), Main Activity (25-30 min), and Wrap-Up/Consolidation (5 min).
        """,
        "Dinleme Aktivitesi Senaryosu": f"""
        Create a listening activity.
        {integration_text}
        - **Crucial Instruction for AI:** The audio script MUST be completely original and unique. Do NOT repeat scripts.
        - Write a short, natural-sounding audio script (100-150 words).
        - After the script, create 5 comprehension questions.
        - End with 'Audio Script' and 'Answer Key' sections.
        """,
        "Ünite Tekrar Testi": f"""
        Create a cumulative unit review test with {kwargs.get('num_questions')} questions.
        - The test must focus on the selected skill: {skill} and topics: {', '.join(kwargs.get('topics', []))}.
        - Include varied question formats.
        - End with an 'Answer Key' section.
        """,
        "Değerlendirme Rubriği": f"""
        Focus Skill: "{skill}"
        {topic_text}
        Create a simple assessment rubric.
        - Use 3 levels: 'Excellent', 'Good', 'Needs Improvement'.
        - Define 3 clear criteria.
        - Provide a brief descriptor for each level.
        """,
        "Ek Çalışma (Farklılaştırılmış)": f"""
        Focus Skill: "{skill}"
        Activity Level: "{kwargs.get('diff_type')}"
        {topic_text}
        {integration_text}
        Create a differentiated activity.
        - If 'Supporting', design a highly-structured task.
        - If 'Expansion', design a task requiring creative or critical thinking.
        - Include an objective and step-by-step instructions.
        """
    }
    
    final_prompt = base_prompt + prompts[tool]
    
    if kwargs.get('clil'):
        final_prompt += "\n- IMPORTANT: Integrate a CLIL component related to technology, science, or digital literacy."
    if kwargs.get('reflection'):
        final_prompt += "\n- IMPORTANT: Add a short, simple reflection question for students to think about their learning."
        
    return final_prompt.strip()


# -----------------------------
# GEMINI API ÇAĞRISI
# -----------------------------
@st.cache_data
def call_gemini_api(_prompt_text):
    try:
        response = model.generate_content(
            _prompt_text,
            generation_config={"max_output_tokens": 2048, "temperature": 0.95}
        )
        return response.text.strip()
    except Exception as e:
        return f"API Hatası: {e}"

# -----------------------------
# ÇIKTI OLUŞTURMA FONKSİYONLARI
# -----------------------------
def create_pdf(content, grade, unit):
    buffer = io.BytesIO()
    margin = 1.27 * cm
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                            rightMargin=margin, leftMargin=margin, 
                            topMargin=margin, bottomMargin=margin)
    
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Heading1_Custom', fontName='DejaVuSans-Bold', 
                              fontSize=11, leading=15, spaceAfter=10))
    styles.add(ParagraphStyle(name='Normal_Custom', fontName='DejaVuSans', 
                              fontSize=10, leading=14, spaceAfter=6))
    styles.add(ParagraphStyle(name='Bullet_Custom', parent=styles['Normal_Custom'], 
                              leftIndent=20))

    Story = []

    def header_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("DejaVuSans", 8)
        canvas.drawString(margin, A4[1] - margin + 0.5*cm, f"{grade} - {unit} | QuickSheet v1.7")
        canvas.drawCentredString(A4[0] / 2.0, 0.75 * cm, f"Sayfa {doc.page}")
        canvas.restoreState()

    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
             Story.append(Spacer(1, 4))
             continue

        formatted_line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)

        if line.startswith('# '):
            Story.append(Paragraph(formatted_line[2:], styles['Heading1_Custom']))
        elif line.startswith('- '):
            p_text = f"• {formatted_line[2:]}"
            p = Paragraph(p_text, styles['Bullet_Custom'])
            Story.append(p)
        else:
            Story.append(Paragraph(formatted_line, styles['Normal_Custom']))

    doc.build(Story, onFirstPage=header_footer, onLaterPages=header_footer)
    buffer.seek(0)
    return buffer

def create_docx(content):
    buffer = io.BytesIO()
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    heading_style = doc.styles.add_style('Heading1_Custom', 1)
    font = heading_style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.bold = True

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text == "":
                continue
            doc.add_paragraph()
            continue

        if line.startswith('# '):
            p = doc.add_paragraph(line[2:], style='Heading1_Custom')
            p.paragraph_format.space_after = Pt(6)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------
# ANA UYGULAMA AKIŞI
# -----------------------------
if st.button("✨ 1. Adım: Materyal Taslağını Oluştur", type="primary", use_container_width=True):
    final_topics = []
    if skill_needed:
        final_topics.extend(selected_topics_from_list)
        if custom_topics:
            final_topics.extend([topic.strip() for topic in custom_topics.split(',') if topic.strip()])
    
    if skill_needed and not final_topics:
        st.warning("Lütfen listeden en az bir alt konu seçin veya kendi konunuzu yazın.")
        st.stop()
    if selected_tool == "Günlük Ders Planı" and not st.session_state.lesson_context:
        st.warning("Lütfen 'Bugünkü Dersin Konusu' alanını doldurun.")
        st.stop()

    prompt_args = {
        "grade": selected_grade, "unit": selected_unit, "skill": selected_skill,
        "topics": final_topics, "clil": include_clil, "reflection": include_reflection,
        "language": plan_language
    }
    if selected_tool in ["Çalışma Sayfası", "Ünite Tekrar Testi"]:
        prompt_args["num_questions"] = num_questions
    if selected_tool == "Ek Çalışma (Farklılaştırılmış)":
        prompt_args["diff_type"] = differentiation_type.split(" ")[0]
    if selected_tool == "Günlük Ders Planı":
        prompt_args["date"] = st.session_state.lesson_date
        prompt_args["context"] = st.session_state.lesson_context

    st.session_state.final_prompt = create_prompt(selected_tool, **prompt_args)
    st.session_state.last_tool = selected_tool
    st.session_state.ai_content = "" 

if 'final_prompt' in st.session_state and st.session_state.final_prompt:
    st.subheader("2. Adım: Komutu Gözden Geçirin (İsteğe Bağlı)")
    edited_prompt = st.text_area("Yapay Zeka Komutu", value=st.session_state.final_prompt, height=250)

    if st.button("🚀 3. Adım: Yapay Zeka ile İçeriği Üret", use_container_width=True):
         with st.spinner(f"{st.session_state.last_tool} üretiliyor..."):
            call_gemini_api.clear()
            st.session_state.ai_content = call_gemini_api(edited_prompt)
            st.session_state.final_prompt = ""

if st.session_state.ai_content:
    st.subheader(f"Üretilen İçerik: {st.session_state.last_tool}")
    edited_content = st.text_area("İçeriği Düzenleyin", value=st.session_state.ai_content, height=400)
    
    st.subheader("Son Adım: İndirin")
    col1, col2 = st.columns(2)
    with col1:
        docx_buffer = create_docx(edited_content)
        st.download_button(
            label="📄 Word Olarak İndir (.docx)", data=docx_buffer,
            file_name=f"{st.session_state.last_tool.replace(' ', '_').lower()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True
        )
    with col2:
        if font_loaded:
            pdf_buffer = create_pdf(edited_content, selected_grade, selected_unit)
            st.download_button(
                label="📑 PDF Olarak İndir", data=pdf_buffer,
                file_name=f"{st.session_state.last_tool.replace(' ', '_').lower()}.pdf",
                mime="application/pdf", use_container_width=True
            )
        else:
            st.warning("Font yüklenemediği için PDF çıktısı devre dışı.")

st.divider()
st.caption("⚡ **Quicksheet v1.7** | Google Gemini API ile güçlendirilmiştir. | MEB 'Türkiye Yüzyılı Maarif Modeli' (2025) 9. Sınıf İngilizce müfredatına uygundur.")
st.caption(f"Geliştirici: Can AKALIN | İletişim: canakalin59@gmail.com | [Instagram](https://instagram.com/can_akalin)")

